#!/usr/bin/env python3
"""
Word报告生成器 - 基于hull_interface数据结构
直接从HullBasicParams、HullCoeffs、HydrostaticData等类生成报告
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
import os
import sys

# 导入hull_interface中的数据结构
from hull_interface import (
    HullBasicParams, HullCoeffs, HydrostaticData, StabilityData, 
    CargoHoldData, ResistanceData, ReportData
)


class WordReportGenerator:
    """基于hull_interface的Word报告生成器"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """设置文档基本格式"""
        # 设置页面
        section = self.doc.sections[0]
        section.page_height = Inches(11.69)  # A4
        section.page_width = Inches(8.27)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # 设置默认字体
        style = self.doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(10.5)
    
    def generate_report_from_data(self, report_data: ReportData, output_path: str):
        """从ReportData生成完整报告"""
        try:
            print("开始生成Word报告...")
            
            self._add_cover_page(report_data)
            self._add_table_of_contents()
            self._add_basic_parameters_section(report_data.basic_params)
            self._add_hull_coefficients_section(report_data.hull_coeffs)
            self._add_hydrostatic_section(report_data.hydro_data)
            self._add_resistance_section(report_data.resistance_data)
            self._add_stability_section(report_data.stability_data)
            self._add_cargo_capacity_section(report_data.cargo_data)
            self._add_summary_section(report_data)
            
            # 保存文档
            self.doc.save(output_path)
            print(f"✅ Word报告生成成功: {output_path}")
            return True
            
        except Exception as e:
            print(f"❌ Word报告生成失败: {str(e)}")
            return False
    
    def _add_cover_page(self, report_data: ReportData):
        """添加封面页"""
        # 主标题
        title = self.doc.add_heading('船舶性能分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(22)
        title_run.bold = True
        
        # 副标题
        subtitle = self.doc.add_heading('船舶设计计算报告', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(16)
        
        self.doc.add_paragraph()  # 空行
        
        # 报告信息表格
        info_table = self.doc.add_table(rows=6, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_table.style = 'Light Grid Accent 1'
        
        # 填充表格数据
        info_data = [
            ('船型', '散货船'),
            ('报告编号', f'RP{datetime.now().strftime("%Y%m%d%H%M")}'),
            ('生成时间', datetime.now().strftime('%Y年%m月%d日 %H:%M')),
            ('计算单位', '船舶性能分析系统'),
            ('主要尺度', f'Lpp={report_data.basic_params.Lpp}m × B={report_data.basic_params.B}m × T={report_data.basic_params.T}m'),
            ('版本', 'v1.0')
        ]
        
        for i, (key, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = key
            info_table.rows[i].cells[1].text = value
            
            # 设置关键单元格样式
            key_cell = info_table.rows[i].cells[0]
            for paragraph in key_cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        self.doc.add_page_break()
    
    def _add_table_of_contents(self):
        """添加目录"""
        heading = self.doc.add_heading('目录', level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(16)
        
        # 目录内容
        toc_items = [
            '1. 船体基础参数',
            '2. 船形系数',
            '3. 静水力性能',
            '4. 阻力与推进性能',
            '5. 稳性分析',
            '6. 舱室容量计算',
            '7. 总结'
        ]
        
        for item in toc_items:
            p = self.doc.add_paragraph(item)
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.line_spacing = 1.5
        
        self.doc.add_page_break()
    
    def _add_basic_parameters_section(self, basic_params: HullBasicParams):
        """添加船体基础参数部分"""
        heading = self.doc.add_heading('1. 船体基础参数', level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        
        # 说明文字
        p = self.doc.add_paragraph('船舶主要几何参数和基础数据，来自用户输入或文件导入。')
        p.paragraph_format.space_after = Pt(12)
        
        # 创建表格
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # 设置表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '参数名称'
        hdr_cells[1].text = '数值'
        hdr_cells[2].text = '单位'
        
        # 设置表头样式
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)
        
        # 从HullBasicParams填充数据
        basic_params_data = [
            ('垂线间长 Lpp', f"{basic_params.Lpp:.2f}", 'm'),
            ('型宽 B', f"{basic_params.B:.2f}", 'm'),
            ('型深 D', f"{basic_params.D:.2f}", 'm'),
            ('吃水 T', f"{basic_params.T:.2f}", 'm'),
            ('排水量 Δ', f"{basic_params.Delta:.1f}", 't')
        ]
        
        # 可选参数
        if basic_params.Loa is not None:
            basic_params_data.append(('总长 Loa', f"{basic_params.Loa:.2f}", 'm'))
        
        basic_params_data.append(('参数来源', basic_params.param_source, '-'))
        
        for param_name, value, unit in basic_params_data:
            row_cells = table.add_row().cells
            row_cells[0].text = param_name
            row_cells[1].text = value
            row_cells[2].text = unit
        
        self.doc.add_paragraph()  # 空行
    
    def _add_hull_coefficients_section(self, hull_coeffs: HullCoeffs):
        """添加船形系数部分"""
        heading = self.doc.add_heading('2. 船形系数', level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        
        p = self.doc.add_paragraph('船形系数是描述船体形状特征的无量纲参数。')
        p.paragraph_format.space_after = Pt(12)
        
        # 创建表格
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # 设置表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '系数名称'
        hdr_cells[1].text = '符号'
        hdr_cells[2].text = '数值'
        
        # 设置表头样式
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)
        
        # 从HullCoeffs填充数据
        coeff_data = [
            ('方形系数', 'Cb', f"{hull_coeffs.Cb:.4f}"),
            ('棱形系数', 'Cp', f"{hull_coeffs.Cp:.4f}"),
            ('水线面系数', 'Cw', f"{hull_coeffs.Cw:.4f}"),
            ('中横剖面系数', 'Cm', f"{hull_coeffs.Cm:.4f}")
        ]
        
        for coeff_name, symbol, value in coeff_data:
            row_cells = table.add_row().cells
            row_cells[0].text = coeff_name
            row_cells[1].text = symbol
            row_cells[2].text = value
        
        # 添加系数说明
        self.doc.add_paragraph()
        explanation_heading = self.doc.add_paragraph('系数说明：')
        explanation_heading.runs[0].bold = True
        
        explanations = [
            f'• 方形系数 Cb：表征船体丰满程度的系数，计算公式：Cb = ∇ / (Lpp × B × T) = {hull_coeffs.Cb:.4f}',
            f'• 棱形系数 Cp：表征船体纵向棱形程度的系数，计算公式：Cp = ∇ / (Am × Lpp) = {hull_coeffs.Cp:.4f}',
            f'• 水线面系数 Cw：表征水线面丰满程度的系数，计算公式：Cw = Aw / (Lpp × B) = {hull_coeffs.Cw:.4f}',
            f'• 中横剖面系数 Cm：表征中横剖面丰满程度的系数，计算公式：Cm = Am / (B × T) = {hull_coeffs.Cm:.4f}'
        ]
        
        for exp in explanations:
            p = self.doc.add_paragraph(exp)
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)
        
        # 添加版本信息
        version_p = self.doc.add_paragraph()
        version_p.add_run(f'系数版本号: {hull_coeffs.coeff_version}').italic = True
    
    def _add_hydrostatic_section(self, hydro_data: HydrostaticData):
        """添加静水力性能部分"""
        heading = self.doc.add_heading('3. 静水力性能', level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        
        p = self.doc.add_paragraph('静水力性能数据基于设计吃水计算，采用梯形积分法。')
        p.paragraph_format.space_after = Pt(12)
        
        # 创建表格
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # 设置表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '参数名称'
        hdr_cells[1].text = '数值'
        hdr_cells[2].text = '单位'
        
        # 设置表头样式
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)
        
        # 从HydrostaticData填充数据
        hydro_data_list = [
            ('排水体积 ▽', f"{hydro_data.displacement_volume:.1f}", 'm³'),
            ('浮心纵向位置 LCB', f"{hydro_data.LCB:.2f}", 'm'),
            ('浮心垂向位置 VCB', f"{hydro_data.VCB:.2f}", 'm'),
            ('每厘米吃水吨数 TPC', f"{hydro_data.TPC:.1f}", 't/cm'),
            ('横稳心半径 BM', f"{hydro_data.BM:.2f}", 'm')
        ]
        
        for param_name, value, unit in hydro_data_list:
            row_cells = table.add_row().cells
            row_cells[0].text = param_name
            row_cells[1].text = value
            row_cells[2].text = unit
        
        # 添加版本信息
        version_p = self.doc.add_paragraph()
        version_p.add_run(f'静水力数据版本号: {hydro_data.hydro_version}').italic = True
    
    def _add_resistance_section(self, resistance_data: ResistanceData):
        """添加阻力与推进性能部分"""
        heading = self.doc.add_heading('4. 阻力与推进性能', level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        
        p = self.doc.add_paragraph(f'基于{resistance_data.method}方法计算的阻力与有效功率。')
        p.paragraph_format.space_after = Pt(12)
        
        if not resistance_data.speed_list:
            self.doc.add_paragraph('无阻力数据可用。')
            return
        
        # 创建表格
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # 设置表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '航速 (kn)'
        hdr_cells[1].text = '总阻力 (kN)'
        hdr_cells[2].text = '有效功率 (kW)'
        hdr_cells[3].text = '备注'
        
        # 设置表头样式
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)
        
        # 从ResistanceData填充数据
        for i in range(len(resistance_data.speed_list)):
            row_cells = table.add_row().cells
            row_cells[0].text = f"{resistance_data.speed_list[i]:.1f}"
            row_cells[1].text = f"{resistance_data.total_resistance_list[i]:.1f}"
            row_cells[2].text = f"{resistance_data.effective_power_list[i]:.1f}"
            row_cells[3].text = '设计航速' if i == len(resistance_data.speed_list) // 2 else ''
    
    def _add_stability_section(self, stability_data: StabilityData):
        """添加稳性分析部分"""
        heading = self.doc.add_heading('5. 稳性分析', level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        
        p = self.doc.add_paragraph(f'基于小角度静稳性假设计算的稳性数据（{stability_data.loading_condition}工况）。')
        p.paragraph_format.space_after = Pt(12)
        
        # 静水力曲线数据表
        sub_heading1 = self.doc.add_heading('静水力曲线数据', level=2)
        sub_heading1_run = sub_heading1.runs[0]
        sub_heading1_run.font.size = Pt(12)
        
        if stability_data.hydro_curve:
            table1 = self.doc.add_table(rows=1, cols=2)
            table1.style = 'Light Grid Accent 1'
            
            # 设置表头
            hdr_cells = table1.rows[0].cells
            hdr_cells[0].text = '吃水 T (m)'
            hdr_cells[1].text = '排水体积 ▽ (m³)'
            
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # 填充静水力曲线数据（显示前10个点）
            for i, (draft, volume) in enumerate(stability_data.hydro_curve[:10]):
                row_cells = table1.add_row().cells
                row_cells[0].text = f"{draft:.2f}"
                row_cells[1].text = f"{volume:.1f}"
            
            if len(stability_data.hydro_curve) > 10:
                self.doc.add_paragraph(f"... 共 {len(stability_data.hydro_curve)} 个数据点")
        else:
            self.doc.add_paragraph('无静水力曲线数据。')
        
        self.doc.add_paragraph()  # 空行
        
        # GZ曲线数据表
        sub_heading2 = self.doc.add_heading('静稳性臂 GZ 曲线数据', level=2)
        sub_heading2_run = sub_heading2.runs[0]
        sub_heading2_run.font.size = Pt(12)
        
        if stability_data.gz_curve:
            table2 = self.doc.add_table(rows=1, cols=2)
            table2.style = 'Light Grid Accent 1'
            
            # 设置表头
            hdr_cells = table2.rows[0].cells
            hdr_cells[0].text = '横倾角 θ (°)'
            hdr_cells[1].text = '稳性臂 GZ (m)'
            
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # 填充GZ曲线数据（显示所有点）
            for angle, gz in stability_data.gz_curve:
                row_cells = table2.add_row().cells
                row_cells[0].text = f"{angle}"
                row_cells[1].text = f"{gz:.3f}"
        else:
            self.doc.add_paragraph('无GZ曲线数据。')
    
    def _add_cargo_capacity_section(self, cargo_data: list[CargoHoldData]):
        """添加舱室容量计算部分"""
        heading = self.doc.add_heading('6. 舱室容量计算', level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        
        p = self.doc.add_paragraph('基于长方体近似法计算的舱室容积。')
        p.paragraph_format.space_after = Pt(12)
        
        if not cargo_data:
            self.doc.add_paragraph('无舱容数据可用。')
            return
        
        # 创建表格
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # 设置表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '舱室名称'
        hdr_cells[1].text = '纵向范围 (m)'
        hdr_cells[2].text = '横向范围 (m)'
        hdr_cells[3].text = '容积 (m³)'
        hdr_cells[4].text = '位置'
        
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # 从CargoHoldData列表填充数据
        total_volume = 0
        for cargo in cargo_data:
            row_cells = table.add_row().cells
            row_cells[0].text = cargo.hold_name
            row_cells[1].text = f"{cargo.hold_position[0]:.1f} - {cargo.hold_position[1]:.1f}"
            row_cells[2].text = f"{cargo.hold_position[2]:.1f} - {cargo.hold_position[3]:.1f}"
            row_cells[3].text = f"{cargo.hold_volume:.1f}"
            row_cells[4].text = self._get_cargo_position_description(cargo.hold_position)
            
            total_volume += cargo.hold_volume
        
        # 添加总计行
        row_cells = table.add_row().cells
        row_cells[0].text = '总计'
        row_cells[1].text = '-'
        row_cells[2].text = '-'
        row_cells[3].text = f"{total_volume:.1f}"
        row_cells[4].text = '-'
        
        # 设置总计行样式
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # 添加总结信息
        self.doc.add_paragraph()
        summary_p = self.doc.add_paragraph()
        summary_p.add_run(f'总舱容: {total_volume:.1f} m³').bold = True
    
    def _get_cargo_position_description(self, position: tuple) -> str:
        """根据舱室位置坐标返回描述"""
        x_start, x_end, y_start, y_end = position
        descriptions = []
        
        # 纵向位置描述
        if x_start >= 0:
            descriptions.append("前部")
        elif x_end <= 0:
            descriptions.append("后部")
        else:
            descriptions.append("中部")
        
        # 横向位置描述
        if y_start >= 0 and y_end >= 0:
            descriptions.append("右舷")
        elif y_start <= 0 and y_end <= 0:
            descriptions.append("左舷")
        else:
            descriptions.append("船中")
        
        return " ".join(descriptions)
    
    def _add_summary_section(self, report_data: ReportData):
        """添加总结部分"""
        heading = self.doc.add_heading('7. 总结', level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        
        # 计算总舱容
        total_cargo = sum(cargo.hold_volume for cargo in report_data.cargo_data) if report_data.cargo_data else 0
        
        summary_text = f"""
本报告基于船舶性能分析系统自动生成，包含了船舶设计的关键性能参数和分析结果。

报告摘要：
• 船舶类型：散货船
• 主要尺度：Lpp = {report_data.basic_params.Lpp:.1f}m × B = {report_data.basic_params.B:.1f}m × T = {report_data.basic_params.T:.1f}m
• 排水量：{report_data.basic_params.Delta:.1f} t
• 方形系数：{report_data.hull_coeffs.Cb:.3f}
• 总舱容：{total_cargo:.1f} m³
• 计算方法：{report_data.resistance_data.method}

所有计算均按照船舶设计规范进行，数据可用于初步设计验证和方案比较。

计算依据：
1. 静水力计算：梯形积分法
2. 阻力估算：{report_data.resistance_data.method}  
3. 稳性分析：小角度静稳性假设
4. 舱容计算：长方体近似法

数据来源：{report_data.basic_params.param_source}

---
报告生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}
生成系统：船舶性能分析系统 v1.0
        """
        
        # 添加总结段落
        paragraphs = summary_text.strip().split('\n')
        for para in paragraphs:
            if para.strip():
                p = self.doc.add_paragraph(para.strip())
                if para.startswith('•'):
                    p.paragraph_format.left_indent = Inches(0.3)
                elif para.startswith('---'):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(12)
